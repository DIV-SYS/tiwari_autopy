from pdf2image import convert_from_path
import pytesseract
import re
from docx import Document
import os

def extract_contact_info(pdf_path):
    images = convert_from_path(pdf_path)
    text = "\n".join(pytesseract.image_to_string(img) for img in images)

    contacts = re.findall(r'([A-Z][a-z]+\s[A-Z][a-z]+).*?([\w\.-]+@[\w\.-]+)', text, re.DOTALL)
    company_match = re.search(r'(Company Name|Owner):\s*(.+)', text)
    company = company_match.group(2).strip() if company_match else "Unknown Company"
    
    result = {
        "company": company,
        "contacts": contacts[:3]
    }
    return result

def save_extracted_data_as_doc(permit_id, info):
    doc = Document()
    doc.add_heading(f'Permit ID: {permit_id}', 0)
    doc.add_paragraph(f"Company: {info['company']}")
    for idx, (name, email) in enumerate(info["contacts"]):
        doc.add_paragraph(f"Contact {idx+1}: {name} - {email}")
    
    path = f"permit_automation/output/processed_files/{permit_id}.docx"
    doc.save(path)
    return path
